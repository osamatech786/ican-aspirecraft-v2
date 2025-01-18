import streamlit as st
from datetime import datetime, date
from streamlit_drawable_canvas import st_canvas
import json
import pandas as pd
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import numpy as np
import smtplib
from email.message import EmailMessage
import re
from dotenv import load_dotenv
import os

# Set page configuration with a favicon
st.set_page_config(
    page_title="AspireCraft Enrolment Form",
    page_icon="https://static.wixstatic.com/media/9699d3_131754d4208c4cfa86e38e9316e5df81~mv2.png/v1/fill/w_204,h_204,al_c,q_85,usm_0.66_1.00_0.01,enc_avif,quality_auto/9699d3_131754d4208c4cfa86e38e9316e5df81~mv2.png",  # Path to your logo
    layout="centered"  # "centered" or "wide"
)

# add render support along with st.secret
def get_secret(key):
    try:
        load_dotenv()
        # Attempt to get the secret from environment variables
        secret = os.environ.get(key)
        if secret is None:
            raise ValueError("Secret not found in environment variables")
        return secret
    except (ValueError, TypeError) as e:
        # If an error occurs, fall back to Streamlit secrets
        if hasattr(st, 'secrets'):
            return st.secrets.get(key)
        # If still not found, return None or handle as needed
        return None

# Load country names and dialing codes from the JSON file
with open("resources/world-countries.json") as file:
    data = json.load(file)
    countries = {entry['name']: entry['dialing_code'] for entry in data}  # Map country name to dialing code
country_names = ["Select"] + sorted(countries.keys())  # Sort country names

# Load nationalities from the JSON file
with open("resources/nationalities.json", "r") as f:
    nationalities = json.load(f)
    
# Load subject areas from the text file
with open("resources/subject_area_list_v2.txt", "r") as file:
    subject_areas = [line.strip() for line in file.readlines()]

# Load and process the CPD Excel file
# df = pd.read_excel('courses.xlsx', sheet_name=0)
df = pd.read_excel('resources/CPD_course_list.xlsx', sheet_name=0)
df = df.drop_duplicates(subset=['Category', 'Course Title'])
category_courses = df.groupby('Category')['Course Title'].apply(list).to_dict()
category_courses = df.groupby('Category')['Course Title'].apply(lambda x: sorted(set(x))).to_dict()

# The sub-options for each subject area, like Foundation, Undergraduate, etc.
sub_options = [
    "Foundation",
    "Undergraduate",
    "Pre-Masters",
    "Postgraduate",
    "PhD & Research",
    "Professional development"
]

# Learning modes (these are hypothetical, but you can adjust them to your case)
learning_modes = ["Online", "Blended", "On-Campus"]

# Function to validate the phone number
def validate_phone_number(phone, dialing_code):
    # Remove all spaces and dashes from the phone number
    phone = phone.replace(" ", "").replace("-", "")
    
    # Check if the phone number starts with the correct dialing code
    # if not phone.startswith(dialing_code):
    #     return False, f"Phone number must start with {dialing_code}."
    
    # Extract the number part (remove the dialing code)
    number_without_code = phone[len(dialing_code):]
    
    # Ensure the number part contains only digits and has a valid length (e.g., 10-15 digits)
    if not number_without_code.isdigit():
        return False, "Phone number must contain only digits after the dialing code."
    
    # if not (10 <= len(number_without_code) <= 15):
    #     return False, "Phone number must be between 10 and 15 digits long (excluding country code)."
    
    return True, ""

def is_valid_email(email):
    # Comprehensive regex for email validation
    pattern = r'''
        ^                         # Start of string
        (?!.*[._%+-]{2})          # No consecutive special characters
        [a-zA-Z0-9._%+-]{1,64}    # Local part: allowed characters and length limit
        (?<![._%+-])              # No special characters at the end of local part
        @                         # "@" symbol
        [a-zA-Z0-9.-]+            # Domain part: allowed characters
        (?<![.-])                 # No special characters at the end of domain
        \.[a-zA-Z]{2,}$           # Top-level domain with minimum 2 characters
    '''
    
    # Match the entire email against the pattern
    return re.match(pattern, email, re.VERBOSE) is not None

def is_signature_drawn(signature):
    # Check if signature is None or an empty numpy array
    if signature is None:
        return False
    # Ensure it is a numpy array and has content
    if isinstance(signature, np.ndarray) and signature.size > 0:
        # Additional check: if the array is not just empty white pixels
        # Assuming white background is [255, 255, 255] in RGB
        if np.all(signature == 255):
            return False
        return True
    return False

# Function to send email with attachments (Handle Local + Uploaded)
def send_email_with_attachments(sender_email, sender_password, receiver_email, subject, body, files=None, local_file_path=None):
    msg = EmailMessage()
    msg['From'] = sender_email
    msg['To'] = ", ".join(receiver_email)
    msg['Subject'] = subject
    msg.set_content(body, subtype='html')

    # Attach uploaded files
    if files:
        for uploaded_file in files:
            uploaded_file.seek(0)  # Move to the beginning of the UploadedFile
            msg.add_attachment(uploaded_file.read(), maintype='application', subtype='octet-stream', filename=uploaded_file.name)

    # Attach local file if specified
    if local_file_path:
        with open(local_file_path, 'rb') as f:
            file_data = f.read()
            file_name = local_file_path.split('/')[-1]
            msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

    # Use Gmail SMTP server for sending the email (office365 for outlook)
    with smtplib.SMTP('smtp.gmail.com', 587) as server:
        server.ehlo()
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)

# Initialize session state variables if they do not exist
if 'step' not in st.session_state:
    st.session_state.step = 1
    st.session_state.submission_done = False
    st.session_state.personal_info = ""  # Full name
    st.session_state.dob = None  # Date of birth
    st.session_state.gender = "Select"  # Gender
    st.session_state.country = ""  # Country of residence
    st.session_state.email = ""  # Email address
    st.session_state.phone = ""  # Phone number
    # st.session_state.address = ""  # Residential address
    st.session_state.current_institution = ""  # Current institution
    # st.session_state.start_date = None  # Uncomment if needed
    st.session_state.front_id_document = None  # Front ID document
    st.session_state.back_id_document = None  # Back ID document
    st.session_state.address_proof = None  # Address proof document
    st.session_state.additional_document = None  # Additional documents if needed
    st.session_state.learning_preferences = ""  # Learning preferences
    st.session_state.special_requirements = ""  # Special requirements
    st.session_state.emergency_contact = ""  # Emergency contact information
    st.session_state.consent = False  # Consent for data processing
    st.session_state.signature = None  # Store signature
    
    st.session_state.sub_option = "Select"  # Default value
    st.session_state.learning_mode = "Select"  # Default value
    st.session_state.qualification_or_experience = "Select"  # Default value
    st.session_state.vocational_sector = "Select"  # Default value
    st.session_state.vocational_other = ""  # Default value
    st.session_state.sector_accreditation = "Select"
    st.session_state.ielts_reason = "Select"
    st.session_state.internship_package = "Select"
    st.session_state.cohort_date = "Select"
    st.session_state.business_services  = None
    st.session_state.category = "Select"  # Default value
    st.session_state.courses = []  # Default to empty list
    st.session_state.learning_mode_cpd = "Online"  # Default value    
    st.session_state.selected_course = {}  # To store selected details    


# Define a function to calculate progress and percentage
def get_progress(step, total_steps=14):
    return int((step / total_steps) * 100)


# Define the total number of steps
total_steps = 14

# Calculate the current progress
progress = get_progress(st.session_state.step, total_steps)

# Display the progress bar and percentage
st.write(f"Progress: {progress}%")
st.progress(progress)


# Define the different steps
if st.session_state.step == 1:
    st.image('resources/AspireCraft_resized.gif', use_column_width=True)
    # st.image(Image.open('resources/logo.png').resize((500, 300)), use_column_width=True)

    st.title("WELCOME TO ASPIRECRAFT!")
    st.write("""
        At AspireCraft, we believe in transforming aspirations into achievements. By joining us, you’re embarking on a journey that unlocks boundless opportunities for growth, success, and excellence. Whether it’s advancing your education, enhancing your skills, or navigating your career.

        AspireCraft is here to guide you every step of the way.
        This form marks the beginning of your story—a story of potential, purpose, and progress. Together, let’s craft your success and empower your future.

        Let’s get started!
    """)
    if st.button("Next"):
        st.session_state.step = 2
        st.experimental_rerun()

elif st.session_state.step == 2:
    st.title("> 1: Personal Information")
    
    # Ensure the personal_info variable is correctly set from the session state
    st.session_state.personal_info = st.text_input(
        "Please enter your full name as it appears on your official documents.",
        value=st.session_state.personal_info  # Retain previous value
    )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.personal_info:  # Check if the field is not empty
            st.session_state.step = 3  # Move to the next step
            st.experimental_rerun()  # Refresh the app to reflect the new step
        else:
            st.warning("Please enter your full name before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 1  # Go back to the previous step (Section 1)
        st.experimental_rerun()  # Refresh to update the step


elif st.session_state.step == 3:
    st.title("> 2: Date of Birth")
    # Check if dob is a string and convert it back to a date object
    if isinstance(st.session_state.get("dob"), str):
        st.session_state.dob = datetime.strptime(st.session_state.get("dob"), "%d-%m-%Y").date()

    # Date of Birth
    # Calculate the maximum allowable date (9 years ago from today)
    max_date = date.today().replace(year=date.today().year - 14)
    
    st.session_state.dob = st.date_input(
        label="Date of Birth",  # Label for the field
        value=st.session_state.get("dob"),  # Correctly access dob from session state
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=max_date,  # Maximum selectable date (at least 9 years old)
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )
    
    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.dob:
            # Convert the selected date to the desired string format (DD-MM-YYYY) only when proceeding to the next step
            # st.session_state.dob = st.session_state.dob.strftime("%d-%m-%Y")

            st.session_state.step = 4
            st.experimental_rerun()
        else:
            st.warning("Please select your date of birth before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 2  # Go back to the previous step (Section 1)
        st.experimental_rerun()


elif st.session_state.step == 4:
    st.title("> 3: Gender")

    # Initialize gender if it doesn't exist
    if 'gender' not in st.session_state:
        st.session_state.gender = "Select"  # Default value

    # Select gender using the selectbox, retaining the previous value
    st.session_state.gender = st.selectbox(
        "Please select your gender.", 
        ["Select", "Male", "Female", "Other"],
        index=["Select", "Male", "Female", "Other"].index(st.session_state.gender)  # Set default value based on session state
    )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.gender != "Select":
            st.session_state.step = 5
            st.experimental_rerun()
        else:
            st.warning("Please select your gender before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 3  # Go back to the previous step (Section 2)
        st.experimental_rerun()


# Step 5: Country selection
elif st.session_state.step == 5:
    st.title("> 4: Country of Residence, Nationality, and Preferred Language")

    # Initialize country, nationality, and preferred language if they don't exist
    if 'country' not in st.session_state:
        st.session_state.country = "Select"  # Default value
    if 'nationality' not in st.session_state:
        st.session_state.nationality = "Select"  # Default value
    if 'preferred_language' not in st.session_state:
        st.session_state.preferred_language = ""  # Default value

    # Select country using the selectbox, retaining the previous value
    st.session_state.country = st.selectbox(
        "Please select your country of residence:", 
        country_names,  # Use the predefined country_names list
        index=country_names.index(st.session_state.country) if st.session_state.country in country_names else 0
    )

    # Select nationality using the selectbox, retaining the previous value
    st.session_state.nationality = st.selectbox(
        "Please select your nationality:",
        ["Select"] + nationalities,  # Load from the JSON file
        index=(["Select"] + nationalities).index(st.session_state.nationality) 
        if st.session_state.nationality in nationalities else 0
    )

    # Text box for Preferred Language of Communication
    st.session_state.preferred_language = st.text_input(
        "Preferred Language of Communication:",
        value=st.session_state.preferred_language,  # Retain the previous input
        placeholder="Enter your preferred language (e.g., English, Spanish)"
    )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.country != "Select" and st.session_state.nationality != "Select" and st.session_state.preferred_language.strip():
            st.session_state.step = 6
            st.experimental_rerun()
        else:
            st.warning("Please complete all fields (Country, Nationality, and Preferred Language) before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 4  # Go back to the previous step (Section 3)
        st.experimental_rerun()

# Step 6: Contact Information
elif st.session_state.step == 6:
    st.title("> 5: Contact Information")

    # Get the selected country's dialing code from the countries dictionary
    selected_dialing_code = countries.get(st.session_state.country, "")

    # Initialize fields if they do not exist
    if 'email' not in st.session_state:
        st.session_state.email = ""  # Default to empty string
    if 'phone' not in st.session_state:
        st.session_state.phone = ""  # Default to empty string
    # if 'address' not in st.session_state:
    #     st.session_state.address = ""  # Default to empty string

    # Input fields for contact information
    st.session_state.email = st.text_input("Please enter your email address where we can reach you.", value=st.session_state.email)

    # Display the country dialing code before the phone number input
    st.session_state.phone = st.text_input(
        f"Please enter your WhatsApp number (international format starting with {selected_dialing_code} for {st.session_state.country}):", 
        value=st.session_state.phone
    )

    # Display the WhatsApp call availability message
    st.markdown(
        """
        ### Ensure WhatsApp Call Availability:
        We may contact you via WhatsApp. Please make sure your phone number is connected to WhatsApp and can receive international calls.
        """
    )
    # Display clickable images in a single line
    st.write("Download WhatsApp for your device:")

    st.markdown(
        """
        <div style="display: flex; justify-content: space-around; align-items: center;">
            <a href="https://play.google.com/store/apps/details?id=com.whatsapp" target="_blank">
                <img src="https://raw.githubusercontent.com/osamatech786/ican-universitysuccess/refs/heads/main/resources/icons/android.png" alt="Download for Android" style="width:100px;height:100px;margin:10px;">
            </a>
            <a href="https://apps.apple.com/app/whatsapp-messenger/id310633997" target="_blank">
                <img src="https://cdn3.iconfinder.com/data/icons/social-media-logos-i-filled-line/2048/5315_-_Apple-512.png" alt="Download for iOS" style="width:100px;height:100px;margin:10px;">
            </a>
            <a href="https://get.microsoft.com/installer/download/9NKSQGP7F2NH" target="_blank">
                <img src="https://github.com/osamatech786/ican-universitysuccess/blob/main/resources/icons/windows.png?raw=true" alt="Download for Windows" style="width:100px;height:100px;margin:10px;">
            </a>
            <a href="https://web.whatsapp.com/desktop/mac_native/release/?configuration=Release" target="_blank">
                <img src="https://github.com/osamatech786/ican-universitysuccess/blob/main/resources/icons/macbook.png?raw=true" alt="Download for Mac" style="width:100px;height:100px;margin:10px;">
            </a>
        </div>
        """, 
        unsafe_allow_html=True
    )

    # Input for address
    # st.session_state.address = st.text_area(
    #     "Please enter your complete mailing address.", 
    #     value=st.session_state.address
    # )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        # if st.session_state.phone and st.session_state.email and st.session_state.address:
        if st.session_state.phone and st.session_state.email:
            if is_valid_email(st.session_state.email):
                is_valid, message = validate_phone_number(st.session_state.phone, selected_dialing_code)
                if is_valid:
                    st.session_state.step = 7
                    st.experimental_rerun()
                else:
                    st.warning(message)
            else:
                st.warning("Please enter a valid email address.")
        else:
            st.warning("Please fill out all the contact information fields before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 5  # Go back to the previous step (Country selection)
        st.experimental_rerun()

elif st.session_state.step == 7:
    st.title("> 6: Educational and Professional Background")

    # Initialize fields if they do not exist
    if 'current_institution' not in st.session_state:
        st.session_state.current_institution = ""  # Default to empty string
    if 'highest_education' not in st.session_state:
        st.session_state.highest_education = "Select"  # Default to Select
    if 'other_education' not in st.session_state:
        st.session_state.other_education = ""  # Default to empty string
    if 'accredited_qualifications' not in st.session_state:
        st.session_state.accredited_qualifications = ""  # Default to empty string
    if 'industry_experience' not in st.session_state:
        st.session_state.industry_experience = "Select"  # Default to Select
    if 'current_role' not in st.session_state:
        st.session_state.current_role = ""  # Default to empty string

    # Input for Current Educational Institution
    st.session_state.current_institution = st.text_input(
        "Please enter the name of your current educational institution (if applicable, else put 'none'):", 
        value=st.session_state.current_institution
    )

    # Dropdown for Highest Level of Education
    st.session_state.highest_education = st.selectbox(
        "Highest Level of Education:",
        [
            "Select",
            "High School Diploma",
            "Bachelor's Degree",
            "Master's Degree",
            "Doctorate",
            "Other (please specify)"
        ],
        index=[
            "Select",
            "High School Diploma",
            "Bachelor's Degree",
            "Master's Degree",
            "Doctorate",
            "Other (please specify)"
        ].index(st.session_state.highest_education)
    )

    # Conditional text box for "Other" education
    if st.session_state.highest_education == "Other (please specify)":
        st.session_state.other_education = st.text_input(
            "Please specify your highest level of education:",
            value=st.session_state.other_education
        )

    # Text field for Accredited Qualifications
    st.session_state.accredited_qualifications = st.text_input(
        "Accredited Qualifications (please specify the sector, if applicable):",
        value=st.session_state.accredited_qualifications
    )

    # Dropdown for Years of Professional Industry Work Experience
    st.session_state.industry_experience = st.selectbox(
        "Years of Professional Industry Work Experience:",
        [
            "Select",
            "Less than 1 year",
            "1–2 years",
            "3–5 years",
            "More than 5 years"
        ],
        index=[
            "Select",
            "Less than 1 year",
            "1–2 years",
            "3–5 years",
            "More than 5 years"
        ].index(st.session_state.industry_experience)
    )

    # Text field for Current Role or Profession
    st.session_state.current_role = st.text_input(
        "Current Role or Profession: (Indicate N/A if not working)",
        value=st.session_state.current_role
    )

    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if (st.session_state.current_institution.strip() and
            st.session_state.highest_education != "Select" and
            (st.session_state.highest_education != "Other (please specify)" or st.session_state.other_education.strip()) and
            st.session_state.industry_experience != "Select" and
            st.session_state.current_role.strip()):
            st.session_state.step = 8
            st.experimental_rerun()
        else:
            st.warning("Please complete all fields before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 6  # Go back to the previous step (Section 5)
        st.experimental_rerun()


# Step 7: Select Services (with Subject Areas, Sub-options, and Learning Modes)
elif st.session_state.step == 8:
    st.title("> 7: Select Services")
        
    # Initialize session state variables
    if "subject_areas" not in st.session_state:
        st.session_state.subject_areas = []
    if "previous_subject_areas" not in st.session_state:
        st.session_state.previous_subject_areas = []

    # Display multiselect for subject areas
    selected_subject_areas = st.multiselect(
        "Please select the subject area(s):",
        sorted(subject_areas),  # Available options
        default=st.session_state.subject_areas  # Retain previous selections
    )

    # Check if the selection has changed
    if selected_subject_areas != st.session_state.previous_subject_areas:
        # Update session state
        st.session_state.subject_areas = selected_subject_areas
        st.session_state.previous_subject_areas = selected_subject_areas
        # Trigger rerun
        st.experimental_rerun()


    # Display relevant options for all selected subject areas
    if st.session_state.subject_areas:
        # st.write("You have selected the following subject area(s):")
        # for area in st.session_state.subject_areas:
        #     st.write(f"- {area}")

        # Handle specific logic for each selected area
        for area in st.session_state.subject_areas:
            
            # Logic for "University Success (Admissions Support)"
            if area == "University Success (Admissions Support)":
                st.subheader(area)
                st.session_state.sub_option = st.selectbox(
                    "Please select your course level.",
                    ["Select"] + sub_options,
                    index=(sub_options.index(st.session_state.sub_option) + 1) if st.session_state.sub_option in sub_options else 0
                )

                st.session_state.learning_mode = st.selectbox(
                    "Please select the learning mode.",
                    ["Select"] + learning_modes,
                    index=(learning_modes.index(st.session_state.learning_mode) + 1) if st.session_state.learning_mode in learning_modes else 0
                )

            # Logic for "International Career Advice and Navigation (ICAN)"
            elif area == "International Career Advice and Navigation (ICAN)":
                st.subheader(area)

                # Career Goals
                st.session_state.career_goals = st.text_area(
                    "Career Goals: Outline your career aspirations and sectors of interest:",
                    value=st.session_state.get("career_goals", "")
                )

                # Reason for Interest
                reason_for_interest_options = [
                    "Select",
                    "Career exploration and development",
                    "CV building and job placement support",
                    "Personalized career diagnostics"
                ]
                st.session_state.reason_for_interest_ican = st.selectbox(
                    "Reason for Interest:",
                    reason_for_interest_options,
                    index=reason_for_interest_options.index(st.session_state.get("reason_for_interest_ican", "Select"))
                )

            # Logic for "Functional Skills Commerce"
            elif area == "Functional Skills Commerce":
                st.subheader(area)

                # Input field for Current Role
                st.session_state.functional_current_role = st.text_input(
                    "Please specify your position (e.g., migrant worker, business owner):",
                    value=st.session_state.get("functional_current_role", "")
                )

                # Select box for Reason for Interest
                functional_reasons = [
                    "Select",
                    "To improve workplace communication and productivity.",
                    "To meet employer expectations for functional skills.",
                    "To access better employment opportunities."
                ]
                st.session_state.functional_reason_for_interest = st.selectbox(
                    "Reason for Interest:",
                    functional_reasons,
                    index=functional_reasons.index(st.session_state.get("functional_reason_for_interest", "Select"))
                )

            # Logic for "Teaching and Assessment Programme"
            elif area == "Teaching and Assessment Programme":
                st.subheader(area)
                # Qualification or experience selection
                st.session_state.qualification_or_experience = st.selectbox(
                    "Please choose one of the following:",
                    ["Select", "An accredited qualification", "More than 2 years of professional industry work experience"],
                    index=["Select", "An accredited qualification", "More than 2 years of professional industry work experience"].index(
                        st.session_state.qualification_or_experience
                    )
                )

                # Display vocational sector options if experience is selected
                if st.session_state.qualification_or_experience == "More than 2 years of professional industry work experience":
                    vocational_options = [
                        "Select",
                        "Health & Social Care",
                        "Construction & Engineering",
                        "Business & Administration",
                        "Digital & IT",
                        "Education & Training",
                        "Retail & Customer Service",
                        "Hospitality & Tourism",
                        "Creative Arts & Media",
                        "Other (please specify)"
                    ]

                    st.session_state.vocational_sector = st.selectbox(
                        "Please select your professional vocational sector:",
                        vocational_options,
                        index=vocational_options.index(st.session_state.vocational_sector) if st.session_state.vocational_sector in vocational_options else 0
                    )

                    # Text input if "Other" is selected
                    if st.session_state.vocational_sector == "Other (please specify)":
                        st.session_state.vocational_other = st.text_input(
                            "Please specify your vocational sector:",
                            value=st.session_state.vocational_other
                        )
                    else:
                        st.session_state.vocational_other = ""
                else:
                    # Reset vocational fields if experience is not selected
                    st.session_state.vocational_sector = "Select"
                    st.session_state.vocational_other = ""

            # Logic for "International Accredited Courses"
            elif area == "International Accredited Courses":
                st.subheader(area)
                accreditation_options = [
                    "Select",
                    "Health, public services and care",
                    "Construction, planning and the built environment",
                    "Information and communication technology",
                    "Arts, media and publishing",
                    "Education and training",
                    "Preparation for Life and Work",
                    "Business, administration and law"
                ]

                # Use concise index calculation
                st.session_state.sector_accreditation = st.selectbox(
                    "Which sector do you wish to achieve accreditation?",
                    accreditation_options,
                    index=accreditation_options.index(st.session_state.sector_accreditation)
                    if st.session_state.sector_accreditation in accreditation_options else 0
                )

            # Logic for "Summer International Internship Programme"
            elif area == "Summer International Internship Programme":
                st.subheader(area)
                # Add a select box for Basic or Premium package selection
                # Internship package selection
                internship_options = ["Select", "Basic", "Premium"]
                st.session_state.internship_package = st.selectbox(
                    "Which package you are interested in? (Basic or Premium)",
                    internship_options,
                    index=internship_options.index(st.session_state.internship_package)
                    if st.session_state.internship_package in internship_options else 0
                )

                # Cohort date selection
                cohort_options = ["Select", "15th July 2025", "29th July 2025", "12th August 2025", "26th August 2025"]
                st.session_state.cohort_date = st.selectbox(
                    "Which cohort date would you like to be part of?",
                    cohort_options,
                    index=cohort_options.index(st.session_state.cohort_date)
                    if st.session_state.cohort_date in cohort_options else 0
                )

                # Reason for Interest selection (single choice)
                reason_options = [
                    "Select",
                    "To gain hands-on work experience in UK hospitals or allied sectors.",
                    "To enhance employability with practical international exposure.",
                    "To participate in cultural and professional development activities."
                ]

                st.session_state.reason_for_interest = st.selectbox(
                    "Please select your reason for interest:",
                    reason_options,
                    index=reason_options.index(st.session_state.reason_for_interest)
                    if st.session_state.get('reason_for_interest') in reason_options else 0
                )

                # Explanations for Basic and Premium packages
                st.write("### Basic Package (Included for All Students):")
                st.markdown(
                    """
                    - **2 weeks work experience** in teaching hospitals across England.
                    - **Accredited UK medical CPD or certification courses** (e.g., ACLS, BLS).
                    - CPD training on patient communication and NHS protocols.
                    - Accommodation in university dorms near teaching hospitals.
                    - Guided London tours to iconic landmarks (e.g., Big Ben, Buckingham Palace).
                    """
                )
                st.write("### Premium Package (Optional Upgrade):")
                st.markdown(
                    """
                    - All Basic Package features.
                    - **Private, upscale accommodation.**
                    - Entertainment: Theme parks, cultural shows, or private river cruises.
                    - **Career mentorship** with UK medical professionals.
                    - Exclusive tours to elite medical facilities (e.g., Royal Society of Medicine).
                    - Concierge services: Airport pickups, personalized itineraries, and meal customization.
                    """
                )

            # Logic for "IELTS Preparation"
            elif area == "IELTS":
                st.subheader(area)
                options = [
                    "Select",
                    "Higher Education Abroad: Admission to universities or colleges in English-speaking countries.",
                    "Immigration: Meeting language requirements for migration to countries like the UK, Canada, or Australia.",
                    "Professional Registration: Certification for professions such as nursing, engineering, or accounting.",
                    "Employment: Enhancing job prospects in international or English-speaking environments.",
                    "Personal Development: Assessing and improving English language proficiency for personal growth."
                ]

                # Use a concise index calculation
                st.session_state.ielts_reason = st.selectbox(
                    "Please select your primary reason for taking the IELTS exam:",
                    options,
                    index=options.index(st.session_state.ielts_reason) if st.session_state.ielts_reason in options else 0
                )

                # Display additional note
                st.write(
                    "_Note: Understanding your motivation helps tailor your preparation effectively._"
                )

            # Logic for "CPD Courses"
            elif area == "CPD Courses":
                st.subheader(area)
                # Category selection
                categories = ["Select"] + list(category_courses.keys())
                st.session_state.category = st.selectbox(
                    "Please select the course category.", 
                    categories, 
                    index=categories.index(st.session_state.category)  # Set default value based on session state
                )

                # Dynamically update course options based on the selected category
                if st.session_state.category != "Select":
                    courses = category_courses.get(st.session_state.category, [])
                else:
                    courses = []

                # Create checkboxes for each course, updating session state appropriately
                selected_courses = st.session_state.courses  # Retrieve previously selected courses

                for course in courses:
                    is_checked = course in selected_courses  # Check if the course is already selected
                    if st.checkbox(course, value=is_checked, key=course):  # Use the course name as the key
                        if course not in selected_courses:
                            selected_courses.append(course)  # Add to the list if checked
                    else:
                        if course in selected_courses:
                            selected_courses.remove(course)  # Remove from the list if unchecked

                # Update session state with selected courses
                st.session_state.courses = selected_courses
                
                st.session_state.learning_mode_cpd = st.selectbox(
                    "Please select your preferred mode of learning.", 
                    ["Online", "Blended", "On-Campus"],
                    index=["Online"].index(st.session_state.learning_mode_cpd)  # Set default based on session state
                )        
                        
            # Logic for "Business Incubation Services"
            elif area == "Business Incubation Services":
                st.subheader(area)

                # Initialize the business_services session state if not already initialized
                if "business_services" not in st.session_state:
                    st.session_state.business_services = []

                # Multi-select for selecting Business Incubation Services services
                selected_services = st.multiselect(
                    "Please select the services you are interested in:",
                    [
                        "MVP Testing: Minimal Viable Product testing and validation.",
                        "Business Plan Review: Expert feedback on your business plan.",
                        "Market Gap Analysis: Identifying opportunities in your target market.",
                        "Workforce Development: Training and support for building your team.",
                        "Growth Management: Strategies for scaling and managing business growth."
                    ],
                    default=st.session_state.business_services  # Retain previous selections
                )

                # Trigger a rerun on change
                if selected_services != st.session_state.business_services:
                    st.session_state.business_services = selected_services
                    st.experimental_rerun()

    else:
        st.write("Please select at least one subject area.")
        
    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        # Ensure at least one subject area is selected
        if not st.session_state.subject_areas:
            st.warning("Please select at least one subject area before proceeding.")
        else:
            # Flag to track overall validation success
            all_valid = True

            # Iterate through each selected subject area
            for subject_area in st.session_state.subject_areas:
                if subject_area == "International Accredited Courses":
                    if st.session_state.sector_accreditation != "Select":
                        st.session_state.selected_course[subject_area] = {
                            'sector_accreditation': st.session_state.sector_accreditation
                        }
                    else:
                        st.warning(f"[{subject_area}] Please select a sector before proceeding.")
                        all_valid = False

                elif subject_area == "IELTS":
                    if st.session_state.ielts_reason != "Select":
                        st.session_state.selected_course[subject_area] = {
                            'ielts_reason': st.session_state.ielts_reason
                        }
                    else:
                        st.warning(f"[{subject_area}] Please select your reason for taking the IELTS exam before proceeding.")
                        all_valid = False

                elif subject_area == "CPD Courses":
                    if st.session_state.courses and st.session_state.learning_mode_cpd != "Select":
                        st.session_state.selected_course[subject_area] = {
                            'category': st.session_state.category,
                            'courses': st.session_state.courses,
                            'learning_mode_cpd': st.session_state.learning_mode_cpd
                        }
                    else:
                        st.warning(f"[{subject_area}] Please select your courses before proceeding.")
                        all_valid = False

                elif subject_area == "International Career Advice and Navigation (ICAN)":
                    if (
                        st.session_state.career_goals.strip() and
                        st.session_state.reason_for_interest_ican != "Select"
                    ):
                        st.session_state.selected_course[subject_area] = {
                            'career_goals': st.session_state.career_goals,
                            'reason_for_interest': st.session_state.reason_for_interest_ican
                        }
                    else:
                        if not st.session_state.career_goals.strip():
                            st.warning(f"[{subject_area}] Please outline your career goals before proceeding.")
                        if st.session_state.reason_for_interest_ican == "Select":
                            st.warning(f"[{subject_area}] Please select your reason for interest before proceeding.")
                        all_valid = False

                elif subject_area == "University Success (Admissions Support)":
                    if st.session_state.sub_option != "Select" and st.session_state.learning_mode != "Select":
                        st.session_state.selected_course[subject_area] = {
                            'course_level': st.session_state.sub_option,
                            'learning_mode': st.session_state.learning_mode
                        }
                    else:
                        st.warning(f"[{subject_area}] Please select the course level and learning mode before proceeding.")
                        all_valid = False

                elif subject_area == "Functional Skills Commerce":
                    if (
                        st.session_state.functional_current_role.strip() and
                        st.session_state.functional_reason_for_interest != "Select"
                    ):
                        st.session_state.selected_course[subject_area] = {
                            "current_role": st.session_state.functional_current_role,
                            "reason_for_interest": st.session_state.functional_reason_for_interest
                        }
                    else:
                        if not st.session_state.functional_current_role.strip():
                            st.warning(f"[{subject_area}] Please specify your current role before proceeding.")
                        if st.session_state.functional_reason_for_interest == "Select":
                            st.warning(f"[{subject_area}] Please select a reason for interest before proceeding.")
                        all_valid = False

                elif subject_area == "Teaching and Assessment Programme":
                    if st.session_state.qualification_or_experience == "Select":
                        st.warning(f"[{subject_area}] Please select an option (accredited qualification or industry experience).")
                        all_valid = False
                    elif (
                        st.session_state.qualification_or_experience == "More than 2 years of professional industry work experience"
                        and (st.session_state.vocational_sector == "Select" or (
                            st.session_state.vocational_sector == "Other (please specify)" and not st.session_state.vocational_other.strip()))
                    ):
                        st.warning(f"[{subject_area}] Please select your vocational sector and specify if applicable.")
                        all_valid = False
                    else:
                        st.session_state.selected_course[subject_area] = {
                            'qualification_or_experience': st.session_state.qualification_or_experience,
                            'vocational_sector': st.session_state.vocational_sector,
                            'vocational_other': st.session_state.vocational_other
                        }

                elif subject_area == "Business Incubation Services":
                    if st.session_state.business_services:
                        st.session_state.selected_course[subject_area] = {
                            'business_services': st.session_state.business_services
                        }
                    else:
                        st.warning(f"[{subject_area}] Please select at least one service to proceed.")
                        all_valid = False

                elif subject_area == "Summer International Internship Programme":
                    if (
                        st.session_state.internship_package != "Select"
                        and st.session_state.cohort_date != "Select"
                        and st.session_state.reason_for_interest != "Select"
                    ):
                        st.session_state.selected_course[subject_area] = {
                            'internship_package': st.session_state.internship_package,
                            'cohort_date': st.session_state.cohort_date,
                            'reason_for_interest': st.session_state.reason_for_interest
                        }
                    else:
                        if st.session_state.internship_package == "Select":
                            st.warning(f"[{subject_area}] Please select a package (Basic or Premium) before proceeding.")
                        if st.session_state.cohort_date == "Select":
                            st.warning(f"[{subject_area}] Please select a cohort date before proceeding.")
                        if st.session_state.reason_for_interest == "Select":
                            st.warning(f"[{subject_area}] Please select a reason for interest before proceeding.")
                        all_valid = False


            # If all validations pass, move to the next step
            if all_valid:
                st.session_state.step = 11
                st.experimental_rerun()
            # else:
            #     st.warning("Please complete all required fields before proceeding.")


    # Handle Back button click
    if back_clicked:
        st.session_state.step = 7  # Go back to the previous step
        st.experimental_rerun()


elif st.session_state.step == 9:
    st.title("> 8: Supporting Documents (Optional)")
    st.write("**Note:** Please note that we will require additional documents to complete enrolment at a later stage.")

    if "files_step_9" not in st.session_state:
        st.session_state.files_step_9 = []

    uploaded_files_9 = st.file_uploader(
        "Upload CV or Portfolio (if applicable):",
        type=["pdf", "docx", "jpg", "png"],
        accept_multiple_files=True,
        key="cv_portfolio"
    )
    if uploaded_files_9:
        for file in uploaded_files_9:
            if file.name not in [f.name for f in st.session_state.files_step_9]:
                st.session_state.files_step_9.append(file)

    if st.session_state.files_step_9:
        st.write("Uploaded CV or Portfolio:")
        for file in st.session_state.files_step_9:
            st.write(f"- {file.name}")

    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    if next_clicked:
        st.session_state.step = 10
        st.experimental_rerun()

    if back_clicked:
        st.session_state.step = 8
        st.experimental_rerun()


elif st.session_state.step == 10:
    st.title("> 9: Supporting Documents (Optional)")
    st.write("**Note:** Please note that we will require additional documents to complete enrolment at a later stage.")

    if "files_step_10" not in st.session_state:
        st.session_state.files_step_10 = []

    uploaded_files_10 = st.file_uploader(
        "Upload Supporting Documents (certificates, qualifications, etc.):",
        type=["jpg", "png", "pdf", "docx"],
        accept_multiple_files=True,
        key="supporting_documents"
    )
    if uploaded_files_10:
        for file in uploaded_files_10:
            if file.name not in [f.name for f in st.session_state.files_step_10]:
                st.session_state.files_step_10.append(file)

    if st.session_state.files_step_10:
        st.write("Uploaded Supporting Documents:")
        for file in st.session_state.files_step_10:
            st.write(f"- {file.name}")

    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    if next_clicked:
        st.session_state.step = 11
        st.experimental_rerun()

    if back_clicked:
        st.session_state.step = 9
        st.experimental_rerun()



elif st.session_state.step == 11:
    st.title("> 10: Additional Information")

    # Initialize fields if they do not exist
    if 'learning_preferences' not in st.session_state:
        st.session_state.learning_preferences = ""  # Default to empty string
    if 'special_requirements' not in st.session_state:
        st.session_state.special_requirements = ""  # Default to empty string
    if 'emergency_contact' not in st.session_state:
        st.session_state.emergency_contact = ""  # Default to empty string
    if 'preferred_start_date' not in st.session_state:
        st.session_state.preferred_start_date = "Select"  # Default to Select
    if 'consent' not in st.session_state:
        st.session_state.consent = False  # Default to unchecked
    # Digital media release consent
    if 'digital_media_consent' not in st.session_state:
        st.session_state.digital_media_consent = True  # Default to unchecked        

    # Input fields with default values from session state
    # Dropdown for Preferred Start Date/Timeline
    st.session_state.preferred_start_date = st.selectbox(
        "Preferred Start Date/Timeline for Participation:",
        [
            "Select",
            "ASAP",
            "1 to 2 months",
            "2 to 4 months",
            "6 months +"
        ],
        index=[
            "Select",
            "ASAP",
            "1 to 2 months",
            "2 to 4 months",
            "6 months +"
        ].index(st.session_state.preferred_start_date) if st.session_state.preferred_start_date in [
            "Select",
            "ASAP",
            "1 to 2 months",
            "2 to 4 months",
            "6 months +"
        ] else 0
    )    
    st.session_state.learning_preferences = st.text_area(
        "Please describe any learning preferences you have.", 
        value=st.session_state.learning_preferences
    )
    st.session_state.special_requirements = st.text_area(
        "Please let us know if you have any special requirements.", 
        value=st.session_state.special_requirements
    )
    st.session_state.emergency_contact = st.text_input(
        "Please provide emergency contact details.", 
        value=st.session_state.emergency_contact
    )

    # Link to the privacy policy
    privacy_policy_doc_link = 'https://drive.google.com/file/d/1sgF6eHZ57idELDEkQD8ZQ7p8VPrmy3WC/view?ts=6789943b'
    st.write(f"[Privacy Policy]({privacy_policy_doc_link})")  # Actual link to privacy policy
    # Privacy policy consent
    st.session_state.consent = st.checkbox(
        "I consent to the collection and processing of my personal data according to AspireCraft’s privacy policy.", 
        value=st.session_state.consent
    )
    
    # Link to the Media Release Consent (M) document
    media_consent_doc_link = 'https://drive.google.com/file/d/1SrHyvp_PHM7OhHyQvE-JJRPCUvloTV0z/view?ts=67899187'
    st.write(f"[Media Release Consent (M) document]({media_consent_doc_link})")  # Actual link to Media Release Consent (M) document
    # Digital media release consent
    st.session_state.digital_media_consent = st.checkbox(
        "I consent to AspireCraft using my photos, videos, or digital media for promotional and educational purposes.", 
        value=st.session_state.digital_media_consent
    )
        
    # Path to the PDF file in the resources folder
    # pdf_file_path = os.path.join('resources', 'Student Privacy Notice_30.07.2024_Rev.1_FF.pdf')
    # Display the link for the PDF file to open in a new tab
    # st.markdown(f'<a href="file://{pdf_file_path}" target="_blank">Privacy Policy</a>', unsafe_allow_html=True)

    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.preferred_start_date != "Select" and all([st.session_state.learning_preferences, st.session_state.special_requirements, st.session_state.emergency_contact, st.session_state.consent]):
            st.session_state.step = 12
            st.experimental_rerun()
        else:
            st.warning("Please complete all fields and consent before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 8
        st.experimental_rerun()


elif st.session_state.step == 12:
    st.title("> 11: Signature")
    st.write("Please provide your signature below:")

    canvas_result = st_canvas(
        stroke_width=2,
        stroke_color="black",
        background_color="white",
        update_streamlit=True,
        height=150,
        width=600,
        drawing_mode="freedraw",
        key="signature_canvas"
    )
    # Only update the session state if there is a change in the canvas
    if canvas_result.image_data is not None:
        st.session_state.signature = canvas_result.image_data

    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if is_signature_drawn(st.session_state.signature):
        # if st.session_state.signature is not None:
            st.session_state.step = 13
            st.experimental_rerun()
        else:
            st.warning("Please provide your signature before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 11  # Go back to the previous step (Section 10)
        st.experimental_rerun()

elif st.session_state.step == 13:
    st.title("Final Review")
    # ==================================================================================================================================
    st.write("Thank you for providing your details. Please review your information.")

    st.header("> Personal Information Review")
   
    st.write(f"**Full Name:** {st.session_state.get('personal_info', 'Not Provided')}")
    st.write(f"**Date of Birth:** {st.session_state.dob.strftime('%d-%m-%Y') if 'dob' in st.session_state else 'Not Provided'}")
    st.write(f"**Gender:** {st.session_state.get('gender', 'Not Provided')}")
    st.write(f"**Country of Residence:** {st.session_state.get('country', 'Not Provided')}")
    st.write(f"**Nationality:** {st.session_state.get('nationality', 'Not Provided')}")
    st.write(f"**Preferred Language of Communication:** {st.session_state.get('preferred_language', 'Not Provided')}")

    st.header("> Contact Information")

    st.write(f"**Email:** {st.session_state.email}")
    st.write(f"**Phone:** {st.session_state.phone}")
    # st.write(f"**Address:** {st.session_state.address}")
    
    st.header("> 6: Educational and Professional Background")

    # Display Current Institution
    st.write(f"**Current Institution:** {st.session_state.current_institution}")
    # Display Highest Level of Education
    if st.session_state.highest_education == "Other (please specify)":
        st.write(f"**Highest Level of Education:** {st.session_state.highest_education} - {st.session_state.other_education}")
    else:
        st.write(f"**Highest Level of Education:** {st.session_state.highest_education}")
    # Display Accredited Qualifications
    st.write(f"**Accredited Qualifications:** {st.session_state.accredited_qualifications}")
    # Display Industry Experience
    st.write(f"**Years of Professional Industry Work Experience:** {st.session_state.industry_experience}")
    # Display Current Role or Profession
    st.write(f"**Current Role or Profession:** {st.session_state.current_role}")

    st.header("> 7: Select Services")

    if "selected_course" in st.session_state and st.session_state.selected_course:
        st.write("**Courses Interested In:**")
        
        for subject_area, details in st.session_state.selected_course.items():
            st.write(f"### {subject_area}")
            
            # Custom logic for each subject area
            if subject_area == "CPD Courses":
                st.write(f"- **Category:** {details.get('category', 'Not Provided')}")
                st.write(f"- **Courses:** {', '.join(details.get('courses', []))}")
                st.write(f"- **Learning Mode:** {details.get('learning_mode_cpd', 'Not Provided')}")

            elif subject_area == "International Career Advice and Navigation (ICAN)":
                st.write(f"- **Career Goals:** {details.get('career_goals', 'Not Provided')}")
                st.write(f"- **Reason for Interest:** {details.get('reason_for_interest', 'Not Provided')}")
            
            elif subject_area == "University Success (Admissions Support)":
                st.write(f"- **Course Level:** {details.get('course_level', 'Not Provided')}")
                st.write(f"- **Learning Mode:** {details.get('learning_mode', 'Not Provided')}")
            
            elif subject_area == "Functional Skills Commerce":
                st.write(f"- **Current Role:** {details.get('current_role', 'Not Provided')}")
                st.write(f"- **Reason for Interest:** {details.get('reason_for_interest', 'Not Provided')}")
            
            elif subject_area == "Teaching and Assessment Programme":
                st.write(f"- **Qualification or Experience:** {details.get('qualification_or_experience', 'Not Provided')}")
                st.write(f"- **Vocational Sector:** {details.get('vocational_sector', 'Not Provided')}")
                if details.get("vocational_other"):
                    st.write(f"- **Other Vocational Sector:** {details.get('vocational_other')}")

            elif subject_area == "International Accredited Courses":
                st.write(f"- **Sector Accreditation:** {details.get('sector_accreditation', 'Not Provided')}")
            
            elif subject_area == "Summer International Internship Programme":
                st.write(f"- **Internship Package:** {details.get('internship_package', 'Not Provided')}")
                st.write(f"- **Cohort Date:** {details.get('cohort_date', 'Not Provided')}")
                st.write(f"- **Reason for Interest:** {details.get('reason_for_interest', 'Not Provided')}")


            elif subject_area == "IELTS":
                st.write(f"- **IELTS Reason:** {details.get('ielts_reason', 'Not Provided')}")

            elif subject_area == "Business Incubation Services":
                st.write(f"- **Business Services:** {', '.join(details.get('business_services', []))}")

            else:
                st.write("No specific fields defined for this subject area.")
    else:
        st.write("**Courses Interested In:** None")



    # st.header("> 8 & 9: Supporting Documents (Optional)")

    # # Combine files from both steps for display purposes
    # all_files = st.session_state.get("files_step_9", []) + st.session_state.get("files_step_10", [])
    # if all_files:
    #     st.write(f"**Total Files Uploaded:** {len(all_files)}")
    #     for file in all_files:
    #         st.write(f"- **File Name:** {file.name}")
    # else:
    #     st.write("No files uploaded.")

        
    st.header("> 10: Additional Information")

    st.write(f"**Preferred Start Date/Timeline for Participation:** {st.session_state.preferred_start_date}")
    st.write(f"**Learning Preferences:** {st.session_state.learning_preferences}")
    st.write(f"**Special Requirements:** {st.session_state.special_requirements}")
    st.write(f"**Emergency Contact:** {st.session_state.emergency_contact}")
    # Display whether digital media consent is provided
    if st.session_state.digital_media_consent:
        st.write("**Digital Media Consent:** Yes, consent provided.")
    else:
        st.write("**Digital Media Consent:** No, consent not provided.")
    

    st.header("> 11: Signature")
    
    if st.session_state.signature is not None:
        st.image(st.session_state.signature, caption="Your Signature")
    

    # Submit button
    submit_clicked = st.button("Submit")

###############################
    with st.spinner('Processing....'):

        # Handle Submit button click
        if submit_clicked:        
            # Create a new Document
            # ======================================================================================================================
            doc = Document()
            doc.add_heading('Enrolment Form Submission', 0)

            # Add Personal Information
            doc.add_heading('Personal Information', level=1)
            doc.add_paragraph(f"Full Name: {st.session_state.get('personal_info', 'Not Provided')}")
            doc.add_paragraph(f"Date of Birth: {st.session_state.dob.strftime('%d-%m-%Y') if 'dob' in st.session_state else 'Not Provided'}")
            doc.add_paragraph(f"Gender: {st.session_state.get('gender', 'Not Provided')}")
            doc.add_paragraph(f"Country of Residence: {st.session_state.get('country', 'Not Provided')}")
            doc.add_paragraph(f"Nationality: {st.session_state.get('nationality', 'Not Provided')}")
            doc.add_paragraph(f"Preferred Language of Communication: {st.session_state.get('preferred_language', 'Not Provided')}")

            # Add Contact Information
            doc.add_heading('Contact Information', level=1)
            doc.add_paragraph(f"Email: {st.session_state.email}")
            doc.add_paragraph(f"Phone: {st.session_state.phone}")
            # doc.add_paragraph(f"Address: {st.session_state.address}")

            # Add Educational and Professional Background
            doc.add_heading('Educational and Professional Background', level=1)
            doc.add_paragraph(f"Current Institution: {st.session_state.current_institution}")
            if st.session_state.highest_education == "Other (please specify)":
                doc.add_paragraph(f"Highest Level of Education: {st.session_state.highest_education} - {st.session_state.other_education}")
            else:
                doc.add_paragraph(f"Highest Level of Education: {st.session_state.highest_education}")
            doc.add_paragraph(f"Accredited Qualifications: {st.session_state.accredited_qualifications}")
            doc.add_paragraph(f"Years of Professional Industry Work Experience: {st.session_state.industry_experience}")
            doc.add_paragraph(f"Current Role or Profession: {st.session_state.current_role}")

            # Add Services and Courses Interested In
            doc.add_heading('Services and Courses Interested In', level=1)
            if "selected_course" in st.session_state and st.session_state.selected_course:
                for subject_area, details in st.session_state.selected_course.items():
                    doc.add_heading(subject_area, level=2)
                    if subject_area == "CPD Courses":
                        doc.add_paragraph(f"Category: {details.get('category', 'Not Provided')}")
                        doc.add_paragraph(f"Courses: {', '.join(details.get('courses', []))}")
                        doc.add_paragraph(f"Learning Mode: {details.get('learning_mode_cpd', 'Not Provided')}")
                    elif subject_area == "International Career Advice and Navigation (ICAN)":
                        doc.add_paragraph(f"Career Goals: {details.get('career_goals', 'Not Provided')}")
                        doc.add_paragraph(f"Reason for Interest: {details.get('reason_for_interest', 'Not Provided')}")
                    elif subject_area == "University Success (Admissions Support)":
                        doc.add_paragraph(f"Course Level: {details.get('course_level', 'Not Provided')}")
                        doc.add_paragraph(f"Learning Mode: {details.get('learning_mode', 'Not Provided')}")
                    elif subject_area == "Functional Skills Commerce":
                        doc.add_paragraph(f"Current Role: {details.get('current_role', 'Not Provided')}")
                        doc.add_paragraph(f"Reason for Interest: {details.get('reason_for_interest', 'Not Provided')}")
                    elif subject_area == "Teaching and Assessment Programme":
                        doc.add_paragraph(f"Qualification or Experience: {details.get('qualification_or_experience', 'Not Provided')}")
                        doc.add_paragraph(f"Vocational Sector: {details.get('vocational_sector', 'Not Provided')}")
                        if details.get("vocational_other"):
                            doc.add_paragraph(f"Other Vocational Sector: {details.get('vocational_other')}")
                    elif subject_area == "International Accredited Courses":
                        doc.add_paragraph(f"Sector Accreditation: {details.get('sector_accreditation', 'Not Provided')}")
                    elif subject_area == "Summer International Internship Programme":
                        doc.add_paragraph(f"Internship Package: {details.get('internship_package', 'Not Provided')}")
                        doc.add_paragraph(f"Cohort Date: {details.get('cohort_date', 'Not Provided')}")
                        doc.add_paragraph(f"Reason for Interest: {details.get('reason_for_interest', 'Not Provided')}")
                    elif subject_area == "IELTS":
                        doc.add_paragraph(f"IELTS Reason: {details.get('ielts_reason', 'Not Provided')}")
                    elif subject_area == "Business Incubation Services":
                        doc.add_paragraph(f"Business Services: {', '.join(details.get('business_services', []))}")
            else:
                doc.add_paragraph("Courses Interested In: None")

            # Add Supporting Documents
            # doc.add_heading('Supporting Documents (Optional)', level=1)
            # all_files = st.session_state.get("files_step_9", []) + st.session_state.get("files_step_10", [])
            # if all_files:
            #     doc.add_paragraph(f"Total Files Uploaded: {len(all_files)}")
            #     for file in all_files:
            #         doc.add_paragraph(f"- File Name: {file.name}")
            # else:
            #     doc.add_paragraph("No files uploaded.")

            # Add Additional Information
            doc.add_heading('Additional Information', level=1)
            doc.add_paragraph(f"Preferred Start Date/Timeline for Participation: {st.session_state.preferred_start_date}")
            doc.add_paragraph(f"Learning Preferences: {st.session_state.learning_preferences}")
            doc.add_paragraph(f"Special Requirements: {st.session_state.special_requirements}")
            doc.add_paragraph(f"Emergency Contact: {st.session_state.emergency_contact}")
            # Add Digital Media Consent to the document
            digital_media_consent_status = "Yes, consent provided." if st.session_state.digital_media_consent else "No, consent not provided."
            doc.add_paragraph(f"Digital Media Consent: {digital_media_consent_status}")
            
            # Add Signature
            doc.add_heading('Signature', level=1)
            if st.session_state.signature is not None:
                # Convert numpy array to PIL image
                image_data = st.session_state.signature
                image = Image.fromarray(image_data.astype(np.uint8))  # Ensure correct data type
                # Save the image to an in-memory file
                image_stream = io.BytesIO()
                image.save(image_stream, format='PNG')
                image_stream.seek(0)
                # Add image to docx
                doc.add_picture(image_stream, width=Inches(2), height=Inches(1))

            # Save the document
            doc_path = f"AspireCraft_Form_Submission_{st.session_state.get('personal_info', 'Unnamed')}.docx"
            doc.save(doc_path)

            # Email
            # ======================================================================================================================
            # Sender email credentials
            # Credentials: Streamlit host st.secrets
            # sender_email = st.secrets["sender_email"]
            # sender_password = st.secrets["sender_password"]

            sender_email = get_secret("sender_email")
            sender_password = get_secret("sender_password")

            # Credentials: Local env
            # load_dotenv()                                     # uncomment import of this library!
            # sender_email = os.getenv('EMAIL')
            # sender_password = os.getenv('PASSWORD')
            team_email = [sender_email]

            learner_email = [st.session_state.email]
            
            subject_team = f"AspireCraft - Country: {st.session_state.country} Name: {st.session_state.personal_info} Submission Date: {date.today()}"
            body_team = "AspireCraft Form submitted. Please find the attached files."

            subject_learner = "Thank You for Signing Up – Next Steps for Your Journey with AspireCraft"
            body_learner = f"""
            <html>
            <body>
                <p>Dear {st.session_state.personal_info},</p>

                <p>Thank you for completing the registration form! We’re excited to have you take the first step toward unlocking your potential with AspireCraft.</p>

                <p><strong>What’s Next?</strong></p>

                <p>We will soon invite you to an <strong>Online Orientation Session</strong>, where you will:</p>
                <ul>
                    <li>Learn more about the program and how it can benefit you.</li>
                    <li>Get detailed guidance on the opportunities available to you.</li>
                    <li>Meet some of our team members who will support you throughout your journey.</li>
                </ul>

                <p><strong>Look Out for Our Invitation</strong></p>
                <p>You’ll receive an email from us shortly with the link to join the orientation session. Please ensure you have a stable internet connection and a device to access the session.</p>

                <p>If you have any questions or need assistance, feel free to contact us:</p>
                <ul>
                    <li><strong>Email:</strong> enquiries@aspirecraft.co.uk</li>
                    <li><strong>WhatsApp:</strong> +44 7711 317561</li>
                </ul>

                <p>We look forward to welcoming you online and helping you craft your success!</p>

                <p><strong>Warm regards,</strong></p>
                <p>The AspireCraft Team<br>
                <em>Crafting Success, Empowering Futures</em></p>
            </body>
            </html>
            """



            # Send email to team with attachments
            if doc_path:
                # send_email_with_attachments(sender_email, sender_password, team_email, subject_team, body_team, all_files, doc_path)
                send_email_with_attachments(sender_email, sender_password, team_email, subject_team, body_team, local_file_path=doc_path)
                # pass
            # Send thank you email to learner
            send_email_with_attachments(sender_email, sender_password, learner_email, subject_learner, body_learner)
            # pass
        
            # Update session state to show the final thank you message
            st.session_state.submission_done = True
            st.session_state.step = 14  # Move to the final step to show the thank you message
            st.experimental_rerun()

#111111111111111111
    # Add a warning before the back button
    st.info("If you go back, you will have to re-sign the form.")

    # Navigation buttons
    back_clicked = st.button("Back", disabled=st.session_state.submission_done)

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 12  # Go back to the previous step
        st.experimental_rerun()
#11111111111111111

# Add a new step for the thank you message
elif st.session_state.step == 14:
    st.title("Thank You!")
    st.write("Check your email for the final boarding.")
    st.write('')
    st.image('resources/AspireCraft.gif', use_column_width=True)

# else:
#     st.write("Form completed. Thank you!")

# Dev : https://linkedin.com/in/osamatech786
